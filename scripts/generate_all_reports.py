import anthropic
import os
import csv
from datetime import datetime

# Word document imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Excel imports
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# PDF imports
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

def analyze_patient_from_csv(patient_row, client):
    """Analyze a patient from CSV data - USED BY ALL THREE FORMATS!"""
    
    patient_summary = f"""
Patient ID: {patient_row['patient_id']}
Name: {patient_row['name']}
Last Appointment: {patient_row['last_appointment']}
Appointments Missed (last 6 months): {patient_row['appointments_missed']}
Medication Adherence: {float(patient_row['medication_adherence'])*100:.0f}%
Crisis Calls (30 days): {patient_row['crisis_calls_30days']}
Diagnosis: {patient_row['diagnosis']}
Case Manager: {patient_row['case_manager']}
"""
    
    prompt = f"""
You are a clinical risk assessment assistant.

Analyze this patient and provide:
1. Risk Level (Low/Medium/High)
2. Primary Risk Factor
3. Recommended Action

PATIENT DATA:
{patient_summary}

Return in this exact format:
Risk Level: [level]
Primary Factor: [factor]
Action: [action]
"""
    
    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=300,
        messages=[{"role": "user", "content": prompt}]
    )
    
    return message.content[0].text

def generate_all_reports(csv_file, timestamp):
    """
    MASTER FUNCTION: Generate Word, Excel, and PDF reports from ONE analysis!
    """
    
    print("=" * 70)
    print("MULTI-FORMAT REPORT GENERATOR")
    print("Generating: Word + Excel + PDF")
    print("=" * 70)
    print()
    
    # Initialize Claude client
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    
    # Read patient data
    print(f"📂 Reading patient data from: {csv_file}")
    patients = []
    
    with open(csv_file, 'r', encoding='utf-8') as file:
        csv_reader = csv.DictReader(file)
        for row in csv_reader:
            patients.append(row)
    
    print(f"✓ Loaded {len(patients)} patients\n")
    
    # ANALYZE ALL PATIENTS ONCE (shared across all formats!)
    print("🤖 Analyzing patients with Claude AI...\n")
    
    patient_analyses = []
    high_risk_patients = []
    medium_risk_patients = []
    low_risk_patients = []
    
    for i, patient in enumerate(patients, 1):
        print(f"   [{i}/{len(patients)}] Analyzing {patient['name']}...")
        
        # Get AI analysis
        analysis = analyze_patient_from_csv(patient, client)
        
        # Parse analysis
        risk_level = "Unknown"
        primary_factor = "Not analyzed"
        action = "No action specified"
        
        for line in analysis.split('\n'):
            if line.startswith('Risk Level:'):
                risk_level = line.replace('Risk Level:', '').strip()
            elif line.startswith('Primary Factor:'):
                primary_factor = line.replace('Primary Factor:', '').strip()
            elif line.startswith('Action:'):
                action = line.replace('Action:', '').strip()
        
        # Store analysis
        patient_analysis = {
            'patient': patient,
            'risk_level': risk_level,
            'primary_factor': primary_factor,
            'action': action
        }
        patient_analyses.append(patient_analysis)
        
        # Categorize by risk
        if "High" in risk_level:
            high_risk_patients.append(patient_analysis)
        elif "Medium" in risk_level:
            medium_risk_patients.append(patient_analysis)
        else:
            low_risk_patients.append(patient_analysis)
    
    print(f"\n✓ Analysis complete!")
    print(f"   High Risk: {len(high_risk_patients)}")
    print(f"   Medium Risk: {len(medium_risk_patients)}")
    print(f"   Low Risk: {len(low_risk_patients)}\n")
    
    # Define output filenames (all with same timestamp!)
    word_file = f"reports/patient_screening_{timestamp}.docx"
    excel_file = f"reports/patient_screening_{timestamp}.xlsx"
    pdf_file = f"reports/patient_screening_{timestamp}.pdf"
    
    print("=" * 70)
    print("GENERATING DOCUMENTS...")
    print("=" * 70)
    print()
    
    # ==================== GENERATE WORD DOCUMENT ====================
    print("📄 Creating Word document...")
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('OAKWOOD BEHAVIORAL HEALTH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Daily Patient Risk Screening Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph()
    date_para.add_run(f'Generated: {datetime.now().strftime("%A, %B %d, %Y at %I:%M %p")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Add patients
    for i, patient_analysis in enumerate(patient_analyses, 1):
        patient = patient_analysis['patient']
        risk_level = patient_analysis['risk_level']
        
        if "High" in risk_level:
            risk_text = "HIGH RISK"
            risk_color = RGBColor(255, 0, 0)
        elif "Medium" in risk_level:
            risk_text = "MEDIUM RISK"
            risk_color = RGBColor(255, 165, 0)
        else:
            risk_text = "LOW RISK"
            risk_color = RGBColor(0, 128, 0)
        
        patient_heading = doc.add_heading(f"{risk_text} - {patient['name']} (ID: {patient['patient_id']})", level=2)
        patient_heading.runs[0].font.color.rgb = risk_color
        
        doc.add_paragraph(f"Case Manager: {patient['case_manager']}")
        doc.add_paragraph(f"Diagnosis: {patient['diagnosis']}")
        doc.add_paragraph(f"Last Appointment: {patient['last_appointment']}")
        doc.add_paragraph(f"Missed Appointments: {patient['appointments_missed']}")
        doc.add_paragraph(f"Medication Adherence: {float(patient['medication_adherence'])*100:.0f}%")
        doc.add_paragraph(f"Crisis Calls (30 days): {patient['crisis_calls_30days']}")
        doc.add_paragraph()
        
        analysis_para = doc.add_paragraph(f"{patient_analysis['risk_level']}\n{patient_analysis['primary_factor']}\n{patient_analysis['action']}")
        analysis_para.runs[0].font.italic = True
        
        doc.add_paragraph('_' * 70)
        doc.add_paragraph()
    
    # Add summary
    doc.add_page_break()
    summary_heading = doc.add_heading('EXECUTIVE SUMMARY', level=1)
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(f"Total Patients Screened: {len(patients)}")
    
    high_para = doc.add_paragraph(f"HIGH RISK: {len(high_risk_patients)} ({len(high_risk_patients)/len(patients)*100:.1f}%)")
    high_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    high_para.runs[0].font.bold = True
    
    medium_para = doc.add_paragraph(f"MEDIUM RISK: {len(medium_risk_patients)} ({len(medium_risk_patients)/len(patients)*100:.1f}%)")
    medium_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)
    medium_para.runs[0].font.bold = True
    
    low_para = doc.add_paragraph(f"LOW RISK: {len(low_risk_patients)} ({len(low_risk_patients)/len(patients)*100:.1f}%)")
    low_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    low_para.runs[0].font.bold = True
    
    doc.save(word_file)
    print(f"   ✓ Word document saved: {word_file}")
    
    # ==================== GENERATE EXCEL SPREADSHEET ====================
    print("📊 Creating Excel spreadsheet...")
    
    wb = Workbook()
    
    # Summary sheet
    ws_summary = wb.active
    ws_summary.title = "Summary"
    
    ws_summary['A1'] = "OAKWOOD BEHAVIORAL HEALTH"
    ws_summary['A1'].font = Font(bold=True, size=16)
    ws_summary.merge_cells('A1:D1')
    
    ws_summary['A2'] = "Daily Patient Risk Screening Report"
    ws_summary['A2'].font = Font(bold=True, size=14)
    ws_summary.merge_cells('A2:D2')
    
    ws_summary['A3'] = f"Generated: {datetime.now().strftime('%A, %B %d, %Y at %I:%M %p')}"
    ws_summary.merge_cells('A3:D3')
    
    # Statistics
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    red_fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
    orange_fill = PatternFill(start_color="FFF4E6", end_color="FFF4E6", fill_type="solid")
    green_fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")
    
    ws_summary['A5'] = "Risk Level"
    ws_summary['B5'] = "Count"
    ws_summary['C5'] = "Percentage"
    ws_summary['D5'] = "Status"
    
    for col in ['A5', 'B5', 'C5', 'D5']:
        ws_summary[col].fill = header_fill
        ws_summary[col].font = Font(bold=True, color="FFFFFF")
    
    ws_summary['A6'] = "HIGH RISK"
    ws_summary['B6'] = len(high_risk_patients)
    ws_summary['C6'] = f"{len(high_risk_patients)/len(patients)*100:.1f}%"
    ws_summary['D6'] = "IMMEDIATE ATTENTION"
    for col in ['A6', 'B6', 'C6', 'D6']:
        ws_summary[col].fill = red_fill
    
    ws_summary['A7'] = "MEDIUM RISK"
    ws_summary['B7'] = len(medium_risk_patients)
    ws_summary['C7'] = f"{len(medium_risk_patients)/len(patients)*100:.1f}%"
    ws_summary['D7'] = "FOLLOW-UP 48-72 HRS"
    for col in ['A7', 'B7', 'C7', 'D7']:
        ws_summary[col].fill = orange_fill
    
    ws_summary['A8'] = "LOW RISK"
    ws_summary['B8'] = len(low_risk_patients)
    ws_summary['C8'] = f"{len(low_risk_patients)/len(patients)*100:.1f}%"
    ws_summary['D8'] = "ROUTINE MONITORING"
    for col in ['A8', 'B8', 'C8', 'D8']:
        ws_summary[col].fill = green_fill
    
    ws_summary.column_dimensions['A'].width = 20
    ws_summary.column_dimensions['B'].width = 12
    ws_summary.column_dimensions['C'].width = 15
    ws_summary.column_dimensions['D'].width = 25
    
    # All Patients sheet
    ws_all = wb.create_sheet("All Patients")
    headers = ["Patient ID", "Name", "Risk Level", "Diagnosis", "Case Manager", 
               "Last Appt", "Missed Appts", "Med Adherence", "Crisis Calls",
               "Primary Risk Factor", "Recommended Action"]
    
    for col_num, header in enumerate(headers, 1):
        cell = ws_all.cell(1, col_num, header)
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
    
    for i, patient_analysis in enumerate(patient_analyses, 2):
        patient = patient_analysis['patient']
        risk_level = patient_analysis['risk_level']
        
        if "High" in risk_level:
            fill = red_fill
        elif "Medium" in risk_level:
            fill = orange_fill
        else:
            fill = green_fill
        
        ws_all.cell(i, 1, patient['patient_id']).fill = fill
        ws_all.cell(i, 2, patient['name']).fill = fill
        ws_all.cell(i, 3, risk_level).fill = fill
        ws_all.cell(i, 4, patient['diagnosis']).fill = fill
        ws_all.cell(i, 5, patient['case_manager']).fill = fill
        ws_all.cell(i, 6, patient['last_appointment']).fill = fill
        ws_all.cell(i, 7, int(patient['appointments_missed'])).fill = fill
        ws_all.cell(i, 8, f"{float(patient['medication_adherence'])*100:.0f}%").fill = fill
        ws_all.cell(i, 9, int(patient['crisis_calls_30days'])).fill = fill
        ws_all.cell(i, 10, patient_analysis['primary_factor']).fill = fill
        ws_all.cell(i, 11, patient_analysis['action']).fill = fill
    
    ws_all.freeze_panes = 'A2'
    
    wb.save(excel_file)
    print(f"   ✓ Excel spreadsheet saved: {excel_file}")
    
    # ==================== GENERATE PDF DOCUMENT ====================
    print("📑 Creating PDF document...")
    
    doc_pdf = SimpleDocTemplate(pdf_file, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'],
                                fontSize=24, textColor=colors.HexColor('#1f4788'),
                                spaceAfter=10, alignment=TA_CENTER, fontName='Helvetica-Bold')
    
    subtitle_style = ParagraphStyle('CustomSubtitle', parent=styles['Heading2'],
                                   fontSize=16, textColor=colors.HexColor('#1f4788'),
                                   spaceAfter=10, alignment=TA_CENTER, fontName='Helvetica-Bold')
    
    # Title page
    elements.append(Spacer(1, 1.5*inch))
    elements.append(Paragraph("OAKWOOD BEHAVIORAL HEALTH", title_style))
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph("Daily Patient Risk Screening Report", subtitle_style))
    elements.append(Spacer(1, 0.2*inch))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%A, %B %d, %Y at %I:%M %p')}", 
                             ParagraphStyle('DateStyle', parent=styles['Normal'], fontSize=12, alignment=TA_CENTER)))
    elements.append(Spacer(1, 0.5*inch))
    
    # Summary table
    summary_data = [
        ['Risk Level', 'Count', 'Percentage', 'Status'],
        ['HIGH RISK', str(len(high_risk_patients)), 
         f"{len(high_risk_patients)/len(patients)*100:.1f}%", 'IMMEDIATE ATTENTION'],
        ['MEDIUM RISK', str(len(medium_risk_patients)), 
         f"{len(medium_risk_patients)/len(patients)*100:.1f}%", 'FOLLOW-UP 48-72 HRS'],
        ['LOW RISK', str(len(low_risk_patients)), 
         f"{len(low_risk_patients)/len(patients)*100:.1f}%", 'ROUTINE MONITORING'],
        ['TOTAL', str(len(patients)), '100%', '']
    ]
    
    summary_table = Table(summary_data, colWidths=[1.5*inch, 1*inch, 1.2*inch, 2*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#ffcccc')),
        ('BACKGROUND', (0, 2), (-1, 2), colors.HexColor('#ffe6cc')),
        ('BACKGROUND', (0, 3), (-1, 3), colors.HexColor('#ccffcc')),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    
    elements.append(summary_table)
    doc_pdf.build(elements)
    
    print(f"   ✓ PDF document saved: {pdf_file}")
    
    print()
    print("=" * 70)
    print("✓ ALL REPORTS GENERATED SUCCESSFULLY!")
    print("=" * 70)
    print()
    
    return {
        'word': word_file,
        'excel': excel_file,
        'pdf': pdf_file,
        'stats': {
            'total': len(patients),
            'high_risk': len(high_risk_patients),
            'medium_risk': len(medium_risk_patients),
            'low_risk': len(low_risk_patients)
        }
    }

# Run the combined generator
if __name__ == "__main__":
    input_csv = "patients.csv"
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    print()
    print("╔" + "═" * 68 + "╗")
    print("║" + " " * 12 + "OAKWOOD BEHAVIORAL HEALTH" + " " * 31 + "║")
    print("║" + " " * 10 + "Multi-Format Report Generator" + " " * 29 + "║")
    print("╚" + "═" * 68 + "╝")
    print()
    
    results = generate_all_reports(input_csv, timestamp)
    
    print("📋 SUMMARY OF GENERATED FILES:")
    print()
    print(f"   📄 Word:  {results['word']}")
    print(f"   📊 Excel: {results['excel']}")
    print(f"   📑 PDF:   {results['pdf']}")
    print()
    print("📊 ANALYSIS STATISTICS:")
    print(f"   Total Patients:   {results['stats']['total']}")
    print(f"   🔴 High Risk:     {results['stats']['high_risk']}")
    print(f"   🟡 Medium Risk:   {results['stats']['medium_risk']}")
    print(f"   🟢 Low Risk:      {results['stats']['low_risk']}")
    print()
    print("=" * 70)
    print("✓ COMPLETE! All files ready in reports/ folder")
    print("=" * 70)
    print()