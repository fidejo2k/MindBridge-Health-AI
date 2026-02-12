import anthropic
import os
import csv
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def analyze_patient_from_csv(patient_row, client):
    """Analyze a patient from CSV data"""
    
    patient_summary = f"""
Patient ID: {patient_row['patient_id']}
Name: {patient_row['name']}
Last Appointment: {patient_row['last_appointment']}
Appointments Missed (last 6 months): {patient_row['appointments_missed']}
Medication Adherence: {float(patient_row['medication_adherence'])*100:.0f}%
Crisis Calls (30 days): {patient_row['crisis_calls_30days']}
Diagnosis: {patient_row['diagnosis']}
Case Manager: {patient_row['case_manager']}
"""
    
    prompt = f"""
You are a clinical risk assessment assistant.

Analyze this patient and provide:
1. Risk Level (Low/Medium/High)
2. Primary Risk Factor
3. Recommended Action

PATIENT DATA:
{patient_summary}

Return in this exact format:
Risk Level: [level]
Primary Factor: [factor]
Action: [action]
"""
    
    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=300,
        messages=[{"role": "user", "content": prompt}]
    )
    
    return message.content[0].text

def create_word_report(csv_file, output_file):
    """Create a professional Word document report"""
    
    # Initialize Claude client
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('OAKWOOD BEHAVIORAL HEALTH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Daily Patient Risk Screening Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date/time
    date_para = doc.add_paragraph()
    date_para.add_run(f'Generated: {datetime.now().strftime("%A, %B %d, %Y at %I:%M %p")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add line break
    doc.add_paragraph()
    
    # Read patient data
    print(f"Reading patient data from: {csv_file}")
    patients = []
    
    with open(csv_file, 'r', encoding='utf-8') as file:
        csv_reader = csv.DictReader(file)
        for row in csv_reader:
            patients.append(row)
    
    print(f"✓ Loaded {len(patients)} patients\n")
    
    # Track statistics
    high_risk_patients = []
    medium_risk_patients = []
    low_risk_patients = []
    
    # Analyze each patient
    print("Analyzing patients and building Word document...\n")
    
    for i, patient in enumerate(patients, 1):
        print(f"[{i}/{len(patients)}] Processing {patient['name']}...")
        
        # Get AI analysis
        analysis = analyze_patient_from_csv(patient, client)
        
        # Determine risk level
        if "High" in analysis:
            risk_level = "HIGH RISK"
            risk_color = RGBColor(255, 0, 0)  # Red
            high_risk_patients.append(patient)
        elif "Medium" in analysis:
            risk_level = "MEDIUM RISK"
            risk_color = RGBColor(255, 165, 0)  # Orange
            medium_risk_patients.append(patient)
        else:
            risk_level = "LOW RISK"
            risk_color = RGBColor(0, 128, 0)  # Green
            low_risk_patients.append(patient)
        
        # Add patient section
        patient_heading = doc.add_heading(f"{risk_level} - {patient['name']} (ID: {patient['patient_id']})", level=2)
        patient_heading_run = patient_heading.runs[0]
        patient_heading_run.font.color.rgb = risk_color
        
        # Add patient details
        doc.add_paragraph(f"Case Manager: {patient['case_manager']}")
        doc.add_paragraph(f"Diagnosis: {patient['diagnosis']}")
        doc.add_paragraph(f"Last Appointment: {patient['last_appointment']}")
        doc.add_paragraph(f"Missed Appointments: {patient['appointments_missed']}")
        doc.add_paragraph(f"Medication Adherence: {float(patient['medication_adherence'])*100:.0f}%")
        doc.add_paragraph(f"Crisis Calls (30 days): {patient['crisis_calls_30days']}")
        
        # Add AI analysis
        doc.add_paragraph()
        analysis_para = doc.add_paragraph(analysis)
        analysis_para.runs[0].font.italic = True
        
        # Add separator
        doc.add_paragraph('_' * 70)
        doc.add_paragraph()
    
    # Add Executive Summary
    doc.add_page_break()
    summary_heading = doc.add_heading('EXECUTIVE SUMMARY', level=1)
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Summary statistics
    doc.add_paragraph(f"Total Patients Screened: {len(patients)}")
    
    high_para = doc.add_paragraph(f"HIGH RISK: {len(high_risk_patients)} ({len(high_risk_patients)/len(patients)*100:.1f}%)")
    high_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    high_para.runs[0].font.bold = True
    
    medium_para = doc.add_paragraph(f"MEDIUM RISK: {len(medium_risk_patients)} ({len(medium_risk_patients)/len(patients)*100:.1f}%)")
    medium_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)
    medium_para.runs[0].font.bold = True
    
    low_para = doc.add_paragraph(f"LOW RISK: {len(low_risk_patients)} ({len(low_risk_patients)/len(patients)*100:.1f}%)")
    low_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    low_para.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # High risk patient list
    if high_risk_patients:
        doc.add_heading('IMMEDIATE ATTENTION REQUIRED:', level=2)
        for patient in high_risk_patients:
            doc.add_paragraph(
                f"• {patient['name']} (ID: {patient['patient_id']}) - Case Manager: {patient['case_manager']}",
                style='List Bullet'
            )
        doc.add_paragraph()
    
    # Medium risk patient list
    if medium_risk_patients:
        doc.add_heading('FOLLOW-UP WITHIN 48-72 HOURS:', level=2)
        for patient in medium_risk_patients:
            doc.add_paragraph(
                f"• {patient['name']} (ID: {patient['patient_id']}) - Case Manager: {patient['case_manager']}",
                style='List Bullet'
            )
    
    # Save document
    doc.save(output_file)
    
    print(f"\n✓ WORD DOCUMENT CREATED!")
    print(f"✓ File saved: {output_file}")
    print(f"\n📊 SUMMARY:")
    print(f"   Total: {len(patients)}")
    print(f"   High Risk: {len(high_risk_patients)}")
    print(f"   Medium Risk: {len(medium_risk_patients)}")
    print(f"   Low Risk: {len(low_risk_patients)}")
    
    return output_file

# Run the script
if __name__ == "__main__":
    input_csv = "patients.csv"
    output_docx = f"reports/patient_screening_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    print("=" * 70)
    print("WORD DOCUMENT GENERATOR")
    print("=" * 70)
    print()
    
    create_word_report(input_csv, output_docx)
    
    print(f"\n📄 Open the Word document: {output_docx}")
    print("\nYou can now EDIT, SHARE, and PRINT this professional report!")
